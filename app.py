import base64
import html as html_module
import io
import os
import random

import streamlit as st
import streamlit.components.v1 as components

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

st.set_page_config(page_title="صانع اختبارات القرآن الكريم الاحترافي", layout="wide")

APP_DIR = os.path.dirname(os.path.abspath(__file__))

# ==================================================================
# قاعدة بيانات القرآن الكريم الموثوقة (نص عثماني - Tanzil Project)
# ==================================================================
# المصدر الوحيد لنص الآيات هو ملف quran-uthmani.txt (نص Tanzil العثماني
# الموثّق، الإصدار 1.1) المرفق مع المشروع. لا يُعدَّل هذا النص برمجياً
# بأي شكل (لا تطبيع، لا حذف تشكيل، لا استبدال حروف) - يُقرأ ويُستخدم
# كما هو تماماً. مسار الملف نسبي لمجلد المشروع حتى يعمل على أي خادم.
QURAN_TEXT_PATH_CANDIDATES = [
    os.path.join(APP_DIR, "data", "quran-uthmani.txt"),
    os.path.join(APP_DIR, "quran-uthmani.txt"),
    os.environ.get("QURAN_TEXT_PATH", ""),
]


class QuranDatabaseError(Exception):
    pass


@st.cache_data(show_spinner=False)
def load_quran_database():
    """
    يقرأ ملف quran-uthmani.txt (تنسيق Tanzil: surah|ayah|text لكل سطر،
    مع تجاهل الأسطر الفارغة وأسطر الترخيص التي تبدأ بـ '#') ويبنيه في
    قاموس {رقم_السورة: {رقم_الآية: النص العثماني كما هو}}.
    """
    path = None
    for candidate in QURAN_TEXT_PATH_CANDIDATES:
        if candidate and os.path.isfile(candidate):
            path = candidate
            break

    if not path:
        raise QuranDatabaseError(
            "تعذّر العثور على ملف نص القرآن الكريم (quran-uthmani.txt). "
            "يجب وضعه داخل مجلد 'data' بجانب app.py."
        )

    quran = {}
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.rstrip("\n\r")
            if not line or line.startswith("#"):
                continue
            parts = line.split("|")
            if len(parts) != 3:
                continue
            s_str, a_str, text = parts
            try:
                s_num, a_num = int(s_str), int(a_str)
            except ValueError:
                continue
            if not text:
                continue
            quran.setdefault(s_num, {})[a_num] = text

    if not quran:
        raise QuranDatabaseError("ملف نص القرآن الكريم موجود لكنه فارغ أو تالف.")

    return quran


try:
    QURAN_TEXT = load_quran_database()
    QURAN_DB_READY = True
    QURAN_DB_ERROR = None
except QuranDatabaseError as e:
    QURAN_TEXT = {}
    QURAN_DB_READY = False
    QURAN_DB_ERROR = str(e)


def get_ayah_text(surah_num, ayah_num):
    """يعيد نص الآية العثماني كما هو من المصدر الموثوق، أو None إن لم توجد."""
    return QURAN_TEXT.get(surah_num, {}).get(ayah_num)


# ==================================================================
# أبعاد الورقة A4 والهوامش - قيم مركزية موحّدة تُستخدم في كل من:
# صفحة الطباعة/المعاينة (HTML) وملف Word (python-docx)
# ==================================================================
PAGE_WIDTH_MM = 210.0
PAGE_HEIGHT_MM = 297.0
PAGE_MARGIN_MM = 12.0  # هامش مخفَّض قليلاً (لا يزال آمناً للطباعة) لإتاحة نموذجين في صفحة A4 واحدة
MODELS_PER_PAGE = 2  # عدد النماذج الكاملة في كل صفحة A4 واحدة

# ==================================================================
# إعدادات الخط العربي المستخدم في معاينة/طباعة HTML (اختياري)
# ==================================================================
# المخرج الأساسي للتطبيق الآن هو ملف Word (.docx) القابل للتعديل، وهو
# لا يحتاج أي ملف خط على الخادم إطلاقاً (يعتمد على خطوط Word نفسها).
# هذا القسم يُستخدم فقط لتحسين شكل معاينة/طباعة HTML داخل المتصفح: إن
# وُجد خط عربي (يُفضّل Amiri) في مجلد fonts سيُضمَّن داخل صفحة
# المعاينة لضمان ثبات الشكل بين الأجهزة، وإلا يعتمد المتصفح على خطوطه
# العربية الخاصة تلقائياً (Tahoma/Arial وغيرها) دون أي عطل.
FONT_DIR_CANDIDATES = [
    os.environ.get("QURAN_EXAM_FONT_DIR", ""),
    os.path.join(APP_DIR, "fonts"),
    APP_DIR,
    os.path.join(APP_DIR, "assets", "fonts"),
    "/usr/share/fonts/truetype/amiri",
    "/usr/share/fonts/opentype/amiri",
    "/usr/share/fonts/truetype/noto",
    "/usr/share/fonts/truetype/dejavu",
    "/usr/share/fonts/truetype/kacst",
    "/usr/share/fonts/truetype/tlwg",
    "/usr/share/fonts",
    "/Library/Fonts",
    "/System/Library/Fonts/Supplemental",
    "C:\\Windows\\Fonts",
]

# ترتيب الأفضلية: يقتصر البحث على خطوط معروف أنها تدعم الرسم العثماني
# القرآني (علامات التشكيل المركّبة، الألف الوصل ٱ، الألف الخنجرية...)
# دعماً كاملاً. لا تُدرَج هنا خطوط عامة مثل DejaVu Sans أو Tahoma أو
# Arial: تغطيتها لرموز يونيكود القرآنية الخاصة غير مكتملة، وتضمينها
# قسراً عبر @font-face يمنع المتصفح من الرجوع تلقائياً إلى خط نظام
# آخر يملك الرمز الناقص - فتظهر مربعات فارغة بدلاً من بعض الحروف.
# إن لم يوجد أي خط من هذه القائمة، يُترك الأمر لخطوط نظام المستخدم
# الخاصة بالمتصفح (والتي عادة ما تتعامل مع هذا التراجع التلقائي بين
# عدة خطوط بشكل أفضل بكثير من خط واحد مُضمَّن قسراً).
FONT_REGULAR_NAMES = [
    "Amiri-Regular.ttf", "Amiri.ttf", "amiri-regular.ttf",
    "NotoNaskhArabic-Regular.ttf", "NotoSansArabic-Regular.ttf",
]
FONT_BOLD_NAMES = [
    "Amiri-Bold.ttf", "amiri-bold.ttf",
    "NotoNaskhArabic-Bold.ttf", "NotoSansArabic-Bold.ttf",
]


def _find_font(candidate_names):
    """يبحث عن أول ملف خط متوفر ضمن مجلدات الخطوط المحتملة."""
    for folder in FONT_DIR_CANDIDATES:
        if not folder or not os.path.isdir(folder):
            continue
        for name in candidate_names:
            path = os.path.join(folder, name)
            if os.path.isfile(path):
                return path
    return None


def get_arabic_font_paths():
    """يعيد (مسار الخط العادي، مسار الخط الغامق أو None، هل هو Amiri؟)."""
    regular = _find_font(FONT_REGULAR_NAMES)
    bold = _find_font(FONT_BOLD_NAMES) or regular
    is_amiri = bool(regular) and "amiri" in os.path.basename(regular).lower()
    return regular, bold, is_amiri


ARABIC_FONT_REGULAR, ARABIC_FONT_BOLD, ARABIC_FONT_IS_AMIRI = get_arabic_font_paths()


def _b64_font(path):
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode("ascii")


def get_embedded_font_css():
    """
    يُضمّن خط العرض/الطباعة داخل HTML كـ base64 (font-face) حتى يظهر
    نفس الشكل تماماً على أي متصفح/جهاز (بدلاً من الاعتماد على خط
    مثبّت على نظام تشغيل المستخدم، والذي قد يختلف بين الأجهزة).
    """
    if not ARABIC_FONT_REGULAR:
        return ""
    try:
        b64_regular = _b64_font(ARABIC_FONT_REGULAR)
    except Exception:
        return ""
    css = f"""
    @font-face {{
        font-family: 'ExamArabicFont';
        src: url(data:font/ttf;base64,{b64_regular}) format('truetype');
        font-weight: normal;
        font-style: normal;
    }}
    """
    if ARABIC_FONT_BOLD and os.path.isfile(ARABIC_FONT_BOLD) and ARABIC_FONT_BOLD != ARABIC_FONT_REGULAR:
        try:
            b64_bold = _b64_font(ARABIC_FONT_BOLD)
            css += f"""
    @font-face {{
        font-family: 'ExamArabicFont';
        src: url(data:font/ttf;base64,{b64_bold}) format('truetype');
        font-weight: bold;
        font-style: normal;
    }}
            """
        except Exception:
            pass
    return css


def esc(text):
    return html_module.escape(str(text if text is not None else ""))
suwar_database = {
    "الفاتحة": {"start": [{"surah_num": 1, "from_ayah": 2, "to_ayah": 5}], "middle": [{"surah_num": 1, "from_ayah": 5, "to_ayah": 7}]},
    "البقرة": {"start": [{"surah_num": 2, "from_ayah": 1, "to_ayah": 5}], "middle": [{"surah_num": 2, "from_ayah": 255, "to_ayah": 255}]},
    "آل عمران": {"start": [{"surah_num": 3, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 3, "from_ayah": 18, "to_ayah": 19}]},
    "النساء": {"start": [{"surah_num": 4, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 4, "from_ayah": 58, "to_ayah": 58}]},
    "المائدة": {"start": [{"surah_num": 5, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 5, "from_ayah": 3, "to_ayah": 3}]},
    "الأنعام": {"start": [{"surah_num": 6, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 6, "from_ayah": 68, "to_ayah": 70}]},
    "الأعراف": {"start": [{"surah_num": 7, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 7, "from_ayah": 204, "to_ayah": 204}]},
    "الأنفال": {"start": [{"surah_num": 8, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 8, "from_ayah": 22, "to_ayah": 22}]},
    "التوبة": {"start": [{"surah_num": 9, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 9, "from_ayah": 128, "to_ayah": 129}]},
    "يونس": {"start": [{"surah_num": 10, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 10, "from_ayah": 62, "to_ayah": 62}]},
    "هود": {"start": [{"surah_num": 11, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 11, "from_ayah": 114, "to_ayah": 114}]},
    "يوسف": {"start": [{"surah_num": 12, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 12, "from_ayah": 92, "to_ayah": 92}]},
    "الرعد": {"start": [{"surah_num": 13, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 13, "from_ayah": 28, "to_ayah": 28}]},
    "إبراهيم": {"start": [{"surah_num": 14, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 14, "from_ayah": 7, "to_ayah": 7}]},
    "الحجر": {"start": [{"surah_num": 15, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 15, "from_ayah": 9, "to_ayah": 9}]},
    "النحل": {"start": [{"surah_num": 16, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 16, "from_ayah": 125, "to_ayah": 125}]},
    "الإسراء": {"start": [{"surah_num": 17, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 17, "from_ayah": 23, "to_ayah": 23}]},
    "الكهف": {"start": [{"surah_num": 18, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 18, "from_ayah": 109, "to_ayah": 109}]},
    "مريم": {"start": [{"surah_num": 19, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 19, "from_ayah": 19, "to_ayah": 19}]},
    "طه": {"start": [{"surah_num": 20, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 20, "from_ayah": 25, "to_ayah": 26}]},
    "الأنبياء": {"start": [{"surah_num": 21, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 21, "from_ayah": 107, "to_ayah": 108}]},
    "الحج": {"start": [{"surah_num": 22, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 22, "from_ayah": 32, "to_ayah": 32}]},
    "المؤمنون": {"start": [{"surah_num": 23, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 23, "from_ayah": 97, "to_ayah": 98}]},
    "النور": {"start": [{"surah_num": 24, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 24, "from_ayah": 35, "to_ayah": 35}]},
    "الفرقان": {"start": [{"surah_num": 25, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 25, "from_ayah": 63, "to_ayah": 63}]},
    "الشعراء": {"start": [{"surah_num": 26, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 26, "from_ayah": 192, "to_ayah": 193}]},
    "النمل": {"start": [{"surah_num": 27, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 27, "from_ayah": 30, "to_ayah": 30}]},
    "القصص": {"start": [{"surah_num": 28, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 28, "from_ayah": 85, "to_ayah": 85}]},
    "العنكبوت": {"start": [{"surah_num": 29, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 29, "from_ayah": 45, "to_ayah": 45}]},
    "الروم": {"start": [{"surah_num": 30, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 30, "from_ayah": 17, "to_ayah": 17}]},
    "لقمان": {"start": [{"surah_num": 31, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 31, "from_ayah": 27, "to_ayah": 27}]},
    "السجدة": {"start": [{"surah_num": 32, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 32, "from_ayah": 16, "to_ayah": 16}]},
    "الأحزاب": {"start": [{"surah_num": 33, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 33, "from_ayah": 56, "to_ayah": 56}]},
    "سبأ": {"start": [{"surah_num": 34, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 34, "from_ayah": 36, "to_ayah": 36}]},
    "فاطر": {"start": [{"surah_num": 35, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 35, "from_ayah": 28, "to_ayah": 28}]},
    "يس": {"start": [{"surah_num": 36, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 36, "from_ayah": 82, "to_ayah": 82}]},
    "الصافات": {"start": [{"surah_num": 37, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 37, "from_ayah": 180, "to_ayah": 182}]},
    "ص": {"start": [{"surah_num": 38, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 38, "from_ayah": 29, "to_ayah": 29}]},
    "الزمر": {"start": [{"surah_num": 39, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 39, "from_ayah": 53, "to_ayah": 53}]},
    "غافر": {"start": [{"surah_num": 40, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 40, "from_ayah": 60, "to_ayah": 60}]},
    "فصلت": {"start": [{"surah_num": 41, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 41, "from_ayah": 30, "to_ayah": 30}]},
    "الشورى": {"start": [{"surah_num": 42, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 42, "from_ayah": 36, "to_ayah": 36}]},
    "الزخرف": {"start": [{"surah_num": 43, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 43, "from_ayah": 88, "to_ayah": 89}]},
    "الدخان": {"start": [{"surah_num": 44, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 44, "from_ayah": 51, "to_ayah": 52}]},
    "الجاثية": {"start": [{"surah_num": 45, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 45, "from_ayah": 20, "to_ayah": 20}]},
    "الأحقاف": {"start": [{"surah_num": 46, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 46, "from_ayah": 35, "to_ayah": 35}]},
    "محمد": {"start": [{"surah_num": 47, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 47, "from_ayah": 24, "to_ayah": 24}]},
    "الفتح": {"start": [{"surah_num": 48, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 48, "from_ayah": 29, "to_ayah": 29}]},
    "الحجرات": {"start": [{"surah_num": 49, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 49, "from_ayah": 13, "to_ayah": 13}]},
    "ق": {"start": [{"surah_num": 50, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 50, "from_ayah": 39, "to_ayah": 39}]},
    "الذاريات": {"start": [{"surah_num": 51, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 51, "from_ayah": 56, "to_ayah": 56}]},
    "الطور": {"start": [{"surah_num": 52, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 52, "from_ayah": 48, "to_ayah": 48}]},
    "النجم": {"start": [{"surah_num": 53, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 53, "from_ayah": 39, "to_ayah": 40}]},
    "القمر": {"start": [{"surah_num": 54, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 54, "from_ayah": 17, "to_ayah": 17}]},
    "الرحمن": {"start": [{"surah_num": 55, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 55, "from_ayah": 26, "to_ayah": 27}]},
    "الواقعة": {"start": [{"surah_num": 56, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 56, "from_ayah": 77, "to_ayah": 78}]},
    "الحديد": {"start": [{"surah_num": 57, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 57, "from_ayah": 16, "to_ayah": 16}]},
    "المجادلة": {"start": [{"surah_num": 58, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 58, "from_ayah": 11, "to_ayah": 11}]},
    "الحشر": {"start": [{"surah_num": 59, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 59, "from_ayah": 21, "to_ayah": 21}]},
    "الممتحنة": {"start": [{"surah_num": 60, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 60, "from_ayah": 5, "to_ayah": 5}]},
    "الصف": {"start": [{"surah_num": 61, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 61, "from_ayah": 10, "to_ayah": 10}]},
    "الجمعة": {"start": [{"surah_num": 62, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 62, "from_ayah": 9, "to_ayah": 9}]},
    "المنافقون": {"start": [{"surah_num": 63, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 63, "from_ayah": 9, "to_ayah": 9}]},
    "التغابن": {"start": [{"surah_num": 64, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 64, "from_ayah": 15, "to_ayah": 15}]},
    "الطلاق": {"start": [{"surah_num": 65, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 65, "from_ayah": 2, "to_ayah": 3}]},
    "التحريم": {"start": [{"surah_num": 66, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 66, "from_ayah": 8, "to_ayah": 8}]},
    "الملك": {"start": [{"surah_num": 67, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 67, "from_ayah": 23, "to_ayah": 23}]},
    "القلم": {"start": [{"surah_num": 68, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 68, "from_ayah": 4, "to_ayah": 5}]},
    "الحاقة": {"start": [{"surah_num": 69, "from_ayah": 2, "to_ayah": 3}], "middle": [{"surah_num": 69, "from_ayah": 48, "to_ayah": 49}]},
    "المعارج": {"start": [{"surah_num": 70, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 70, "from_ayah": 19, "to_ayah": 20}]},
    "نوح": {"start": [{"surah_num": 71, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 71, "from_ayah": 10, "to_ayah": 11}]},
    "الجن": {"start": [{"surah_num": 72, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 72, "from_ayah": 13, "to_ayah": 13}]},
    "المزمل": {"start": [{"surah_num": 73, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 73, "from_ayah": 20, "to_ayah": 20}]},
    "المدثر": {"start": [{"surah_num": 74, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 74, "from_ayah": 38, "to_ayah": 39}]},
    "القيامة": {"start": [{"surah_num": 75, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 75, "from_ayah": 16, "to_ayah": 17}]},
    "الإنسان": {"start": [{"surah_num": 76, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 76, "from_ayah": 22, "to_ayah": 22}]},
    "المرسلات": {"start": [{"surah_num": 77, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 77, "from_ayah": 41, "to_ayah": 42}]},
    "النبأ": {"start": [{"surah_num": 78, "from_ayah": 1, "to_ayah": 11}], "middle": [{"surah_num": 78, "from_ayah": 21, "to_ayah": 26}]},
    "النازعات": {"start": [{"surah_num": 79, "from_ayah": 1, "to_ayah": 10}], "middle": [{"surah_num": 79, "from_ayah": 15, "to_ayah": 20}]},
    "عبس": {"start": [{"surah_num": 80, "from_ayah": 1, "to_ayah": 13}], "middle": [{"surah_num": 80, "from_ayah": 17, "to_ayah": 22}]},
    "التكوير": {"start": [{"surah_num": 81, "from_ayah": 1, "to_ayah": 7}], "middle": [{"surah_num": 81, "from_ayah": 15, "to_ayah": 26}]},
    "الانفطار": {"start": [{"surah_num": 82, "from_ayah": 1, "to_ayah": 4}], "middle": [{"surah_num": 82, "from_ayah": 13, "to_ayah": 19}]},
    "المطففين": {"start": [{"surah_num": 83, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 83, "from_ayah": 22, "to_ayah": 23}]},
    "الانشقاق": {"start": [{"surah_num": 84, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 84, "from_ayah": 7, "to_ayah": 8}]},
    "البروج": {"start": [{"surah_num": 85, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 85, "from_ayah": 12, "to_ayah": 13}]},
    "الطارق": {"start": [{"surah_num": 86, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 86, "from_ayah": 13, "to_ayah": 14}]},
    "الأعلى": {"start": [{"surah_num": 87, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 87, "from_ayah": 14, "to_ayah": 15}]},
    "الغاشية": {"start": [{"surah_num": 88, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 88, "from_ayah": 21, "to_ayah": 22}]},
    "الفجر": {"start": [{"surah_num": 89, "from_ayah": 2, "to_ayah": 3}], "middle": [{"surah_num": 89, "from_ayah": 27, "to_ayah": 28}]},
    "البلد": {"start": [{"surah_num": 90, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 90, "from_ayah": 8, "to_ayah": 9}]},
    "الشمس": {"start": [{"surah_num": 91, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 91, "from_ayah": 9, "to_ayah": 10}]},
    "الليل": {"start": [{"surah_num": 92, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 92, "from_ayah": 5, "to_ayah": 6}]},
    "الضحى": {"start": [{"surah_num": 93, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 93, "from_ayah": 5, "to_ayah": 6}]},
    "الشرح": {"start": [{"surah_num": 94, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 94, "from_ayah": 5, "to_ayah": 7}]},
    "التين": {"start": [{"surah_num": 95, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 95, "from_ayah": 4, "to_ayah": 5}]},
    "العلق": {"start": [{"surah_num": 96, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 96, "from_ayah": 6, "to_ayah": 7}]},
    "القدر": {"start": [{"surah_num": 97, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 97, "from_ayah": 3, "to_ayah": 4}]},
    "البينة": {"start": [{"surah_num": 98, "from_ayah": 1, "to_ayah": 1}], "middle": [{"surah_num": 98, "from_ayah": 7, "to_ayah": 7}]},
    "الزلزلة": {"start": [{"surah_num": 99, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 99, "from_ayah": 7, "to_ayah": 8}]},
    "العاديات": {"start": [{"surah_num": 100, "from_ayah": 1, "to_ayah": 3}], "middle": [{"surah_num": 100, "from_ayah": 6, "to_ayah": 7}]},
    "القارعة": {"start": [{"surah_num": 101, "from_ayah": 2, "to_ayah": 3}], "middle": [{"surah_num": 101, "from_ayah": 6, "to_ayah": 7}]},
    "التكاثر": {"start": [{"surah_num": 102, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 102, "from_ayah": 5, "to_ayah": 6}]},
    "العصر": {"start": [{"surah_num": 103, "from_ayah": 2, "to_ayah": 3}], "middle": [{"surah_num": 103, "from_ayah": 3, "to_ayah": 3}]},
    "الهمزة": {"start": [{"surah_num": 104, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 104, "from_ayah": 4, "to_ayah": 6}]},
    "الفيل": {"start": [{"surah_num": 105, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 105, "from_ayah": 3, "to_ayah": 5}]},
    "قريش": {"start": [{"surah_num": 106, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 106, "from_ayah": 3, "to_ayah": 4}]},
    "الماعون": {"start": [{"surah_num": 107, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 107, "from_ayah": 5, "to_ayah": 7}]},
    "الكوثر": {"start": [{"surah_num": 108, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 108, "from_ayah": 1, "to_ayah": 3}]},
    "الكافرون": {"start": [{"surah_num": 109, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 109, "from_ayah": 5, "to_ayah": 6}]},
    "النصر": {"start": [{"surah_num": 110, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 110, "from_ayah": 3, "to_ayah": 3}]},
    "المسد": {"start": [{"surah_num": 111, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 111, "from_ayah": 3, "to_ayah": 5}]},
    "الإخلاص": {"start": [{"surah_num": 112, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 112, "from_ayah": 3, "to_ayah": 4}]},
    "الفلق": {"start": [{"surah_num": 113, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 113, "from_ayah": 3, "to_ayah": 5}]},
    "الناس": {"start": [{"surah_num": 114, "from_ayah": 1, "to_ayah": 2}], "middle": [{"surah_num": 114, "from_ayah": 4, "to_ayah": 6}]},
}

# ==================================================================
# محرك توليد أسئلة القرآن الكريم
# ==================================================================
# يبني "مسبح" (pool) من الأسئلة الفريدة لكل سورة مختارة (بداية / وسط)،
# ثم يختار منه بدون تكرار داخل نفس النموذج، ويحاول تقليل التشابه بين
# النماذج المختلفة إلى سؤال مشترك واحد كحد أقصى، دون الدخول أبداً في
# حلقة لا نهائية (يوجد سقف محاولات صريح في كل مكان).

MAX_SELECTION_ATTEMPTS = 300  # سقف صريح لمنع أي حلقة لا نهائية


def build_question_pool(selected_suwar, suwar_database, category):
    """
    يبني قائمة أسئلة فريدة (سورة + رقم آية البداية + رقم آية النهاية)
    لفئة معيّنة (start/middle). هوية السؤال (id) تعتمد على أرقام الآيات
    الفعلية (وليس على الفئة) - حتى يُكتشف التكرار بشكل صحيح إن ظهر نفس
    المدى بالضبط في كلا القسمين. النص الفعلي يُجلب مباشرة من قاعدة
    بيانات القرآن الموثوقة (quran-uthmani.txt) عبر get_ayah_text - لا
    يوجد أي نص قرآني مخزَّن مباشرة هنا.
    """
    pool = []
    for sura in selected_suwar:
        entries = suwar_database.get(sura, {}).get(category, [])
        for entry in entries:
            surah_num = entry["surah_num"]
            from_ayah = entry["from_ayah"]
            to_ayah = entry["to_ayah"]
            from_text = get_ayah_text(surah_num, from_ayah)
            to_text = get_ayah_text(surah_num, to_ayah)
            if from_text is None or to_text is None:
                # لن يحدث في البيانات المُتحقَّق منها، لكن نتجاهل بأمان
                # بدلاً من التعطّل إن حدث خطأ غير متوقع في المصدر.
                continue
            pool.append({
                "id": (sura, from_ayah, to_ayah),
                "surah": sura,
                "category": category,
                "surah_num": surah_num,
                "from_ayah": from_ayah,
                "to_ayah": to_ayah,
                "from": from_text,
                "to": to_text,
            })
    return pool


def _sample_unique(pool, count, rng):
    """يختار (بدون تكرار) count عنصر من pool. يعيد أقل عدد متاح إن لم يكفِ."""
    if count <= 0 or not pool:
        return []
    n = min(count, len(pool))
    return rng.sample(pool, n)


def select_questions_for_model(start_pool, middle_pool, q_start_count, q_middle_count, rng):
    """
    يختار أسئلة نموذج واحد (بداية ثم وسط) مع ترقيم متواصل لاحقاً.
    يمنع التكرار داخل نفس النموذج حتى لو كان نفس السؤال (نفس السورة
    ونفس نص "من"/"إلى") موجوداً في كل من قسمي "البداية" و"الوسط" معاً -
    عبر استبعاد هويات أسئلة البداية المختارة من مسبح أسئلة الوسط قبل
    الاختيار منه.
    """
    warnings = []

    chosen_start = _sample_unique(start_pool, q_start_count, rng)
    if len(chosen_start) < q_start_count:
        warnings.append(
            f"عدد أسئلة (البداية) المتاحة في السور المختارة هو {len(start_pool)} فقط، "
            f"وقد طُلب {q_start_count}. تم استخدام الحد الأقصى المتاح ({len(chosen_start)})."
        )

    used_ids = _question_ids(chosen_start)
    middle_pool_available = [q for q in middle_pool if q["id"] not in used_ids]

    chosen_middle = _sample_unique(middle_pool_available, q_middle_count, rng)
    if len(chosen_middle) < q_middle_count:
        warnings.append(
            f"عدد أسئلة (الوسط) المتاحة (غير المكرّرة مع أسئلة البداية) في السور "
            f"المختارة هو {len(middle_pool_available)} فقط، وقد طُلب {q_middle_count}. "
            f"تم استخدام الحد الأقصى المتاح ({len(chosen_middle)})."
        )

    return chosen_start + chosen_middle, warnings


def _question_ids(questions):
    return frozenset(q["id"] for q in questions)


def generate_all_models(num_models, q_start_count, q_middle_count, selected_suwar,
                         suwar_database, seed=None):
    """
    يولّد عدد num_models من النماذج، بحيث:
      - لا يوجد سؤال مكرر داخل نفس النموذج.
      - لا يشترك أي نموذجين في أكثر من سؤال واحد (قدر الإمكان).
      - لا يحدث أي تعليق/حلقة لا نهائية مهما كان عدد الأسئلة المتاحة.
    يعيد (models, general_warnings) حيث models قائمة من:
      {"questions": [...], "warnings": [...]}
    """
    start_pool = build_question_pool(selected_suwar, suwar_database, "start")
    middle_pool = build_question_pool(selected_suwar, suwar_database, "middle")

    general_warnings = []
    if not selected_suwar:
        general_warnings.append("لم يتم اختيار أي سورة، لذلك لا يمكن توليد أسئلة القرآن الكريم.")

    base_seed = seed if seed is not None else random.randint(1, 10_000_000)

    models = []
    existing_id_sets = []

    for model_index in range(num_models):
        best_choice = None
        best_overlap = None

        for attempt in range(MAX_SELECTION_ATTEMPTS):
            trial_rng = random.Random(base_seed * 1_000_003 + model_index * 97 + attempt)
            questions, warnings = select_questions_for_model(
                start_pool, middle_pool, q_start_count, q_middle_count, trial_rng
            )
            ids = _question_ids(questions)

            if not existing_id_sets:
                best_choice = (questions, warnings)
                best_overlap = 0
                break

            max_overlap = max(len(ids & prev_ids) for prev_ids in existing_id_sets)

            if best_overlap is None or max_overlap < best_overlap:
                best_choice = (questions, warnings)
                best_overlap = max_overlap

            if max_overlap <= 1:
                break

        questions, warnings = best_choice
        if best_overlap is not None and best_overlap > 1:
            warnings = warnings + [
                f"النموذج {model_index + 1}: عدد الأسئلة المتاحة في السور المختارة قليل جداً، "
                f"فتعذّر تقليل الأسئلة المشتركة مع نموذج سابق إلى سؤال واحد فقط. "
                f"يُنصح باختيار سور إضافية لزيادة التنوع."
            ]

        existing_id_sets.append(_question_ids(questions))
        models.append({"questions": questions, "warnings": warnings})

    return models, general_warnings
# ==================================================================
# مواصفة صفحة الاختبار (Single Source of Truth)
# ==================================================================
# هذه القائمة من "الكتل" (blocks) هي المصدر الوحيد لمحتوى وترتيب
# صفحة الاختبار. كل من: معاينة Streamlit، صفحة الطباعة (HTML)،
# وملف Word تقرأ نفس القائمة بالضبط ولا يُعاد توليد أي محتوى بشكل
# منفصل في أي منها.

def question_line_text(q):
    # عندما يكون المدى آية واحدة فقط (من == إلى)، تُعرض عبارة واحدة
    # بنص الآية الموثوق كاملاً بدلاً من تكرار نفس النص في قالب
    # "من قوله ... إلى قوله" (والذي يصبح غير منطقي حين تكون البداية
    # والنهاية نفس الآية).
    if q.get("from_ayah") is not None and q.get("to_ayah") is not None and q["from_ayah"] == q["to_ayah"]:
        return (
            f"اقرأ مستعيناً بالله تعالى من سورة {q['surah']} قوله تعالى: "
            f"[ {q['from']} ]."
        )
    return (
        f"اقرأ مستعيناً بالله تعالى من سورة {q['surah']} من قوله تعالى: "
        f"[ {q['from']} ] إلى قوله تعالى: [ {q['to']} ]."
    )


ORDINAL_WORDS = ["أولاً", "ثانياً", "ثالثاً", "رابعاً", "خامساً", "سادساً"]


def build_exam_page_blocks(exam_type, grade_option, level_option, model_index,
                            questions, include_tarbawy, include_tafsir,
                            custom_tafsir_verse, include_meaning):
    """يبني قائمة الكتل المكوّنة لصفحة اختبار واحدة (نموذج واحد)."""
    blocks = []

    blocks.append({
        "type": "header_box",
        "exam_title": f"امتحان {exam_type} للعام القرآني",
        "date_note": "التاريخ: ............ ١٤هـ  -  ............ ٢٠م",
        "info_fields": [
            f"الفرقة: {grade_option}",
            f"المستوى: {level_option}",
            "رقم الحلقة: (      )",
        ],
        "office_field": "مكتب: ....................................................",
        "name_field": "اسم الطالب: ...................................................................",
    })

    blocks.append({"type": "model_title", "text": f"نموذج رقم ({model_index})"})

    # -- ترقيم الأقسام ديناميكياً حسب المحاور المفعّلة فقط (بدون أرقام
    # ثابتة/مُدرَجة يدوياً) - أسئلة القرآن الكريم دائماً أولاً -------
    enabled_sections = ["quran"]
    if include_tarbawy:
        enabled_sections.append("tarbawy")
    if include_tafsir:
        enabled_sections.append("tafsir")
    if include_meaning:
        enabled_sections.append("meaning")

    ordinal_of = {}
    for i, key in enumerate(enabled_sections):
        ordinal_of[key] = ORDINAL_WORDS[i] if i < len(ORDINAL_WORDS) else str(i + 1)

    if questions:
        blocks.append({"type": "section_heading", "text": f"{ordinal_of['quran']}: أسئلة القرآن الكريم:", "color": "#1E3A8A"})
        for q in questions:
            blocks.append({"type": "question", "number": q["number"], "text": question_line_text(q)})
    else:
        blocks.append({"type": "note", "text": f"{ordinal_of['quran']}: أسئلة القرآن الكريم: (لم يتم اختيار أسئلة)"})

    if include_tarbawy:
        blocks.append({"type": "section_heading", "text": f"{ordinal_of['tarbawy']}: التربوي"})
        blocks.append({"type": "dotted_line", "label": "فقه:"})
        blocks.append({"type": "dotted_line", "label": "عقيدة:"})

    if include_tafsir:
        blocks.append({"type": "section_heading", "text": f"{ordinal_of['tafsir']}: التفسير الإجمالي"})
        # لا يُدرَج أي نص أو آية تلقائياً؛ إن أدخل المستخدم نصاً في الشريط
        # الجانبي يظهر بعد "س:"، وإلا تبقى المساحة فارغة للكتابة اليدوية.
        tafsir_text = (custom_tafsir_verse or "").strip()
        blocks.append({"type": "dotted_line", "label": "س:", "prefix_text": tafsir_text})
        blocks.append({"type": "dotted_line", "label": "جـ:"})

    if include_meaning:
        blocks.append({"type": "section_heading", "text": f"{ordinal_of['meaning']}: معاني الكلمات:"})
        blocks.append({"type": "meaning_grid", "cells": [
            "١- " + "." * 28, "٢- " + "." * 28,
            "٣- " + "." * 28, "٤- " + "." * 28,
            "٥- " + "." * 60, "",
        ]})

    return blocks


def build_exam_model_data(model_index, exam_type, grade_option, level_option,
                           questions, include_tarbawy, include_tafsir,
                           custom_tafsir_verse, include_meaning, warnings):
    """
    يبني كائن النموذج الكامل: بيانات وصفية (للاستخدام في تسمية الملفات
    والتحذيرات) + قائمة blocks الموحّدة التي يقرأها كل مُصدِّر.
    """
    numbered_questions = []
    for i, q in enumerate(questions, start=1):
        numbered_questions.append({
            "number": i, "surah": q["surah"], "from": q["from"], "to": q["to"],
            "from_ayah": q.get("from_ayah"), "to_ayah": q.get("to_ayah"),
        })

    blocks = build_exam_page_blocks(
        exam_type=exam_type,
        grade_option=grade_option,
        level_option=level_option,
        model_index=model_index,
        questions=numbered_questions,
        include_tarbawy=include_tarbawy,
        include_tafsir=include_tafsir,
        custom_tafsir_verse=custom_tafsir_verse,
        include_meaning=include_meaning,
    )

    return {
        "model_index": model_index,
        "exam_type": exam_type,
        "grade_option": grade_option,
        "level_option": level_option,
        "questions": numbered_questions,
        "include_tarbawy": include_tarbawy,
        "include_tafsir": include_tafsir,
        "tafsir_verse": custom_tafsir_verse,
        "include_meaning": include_meaning,
        "warnings": warnings,
        "blocks": blocks,
    }
# ==================================================================
# مُصيِّر HTML/طباعة — التصميم المرجعي الوحيد لصفحة الاختبار
# ==================================================================
# هذا المُصيِّر يقرأ نفس قائمة blocks المُستخدَمة في ملف Word، وهو ما
# يظهر داخل معاينة Streamlit وأيضاً في صفحة الطباعة النظيفة (بدون
# الشريط الجانبي أو أي عناصر واجهة أخرى). يمكن أيضاً طباعة هذه الصفحة
# أو حفظها كـ PDF مباشرة من المتصفح/نظام التشغيل عند الحاجة.

DOTS_LONG = "." * 130


def render_blocks_to_html(blocks):
    parts = []
    for b in blocks:
        t = b["type"]
        if t == "header_box":
            info_spans = "".join(f'<span class="info-field">{esc(f)}</span>' for f in b["info_fields"])
            parts.append(
                '<div class="header-box">'
                '<div class="header-title-row">'
                f'<span class="header-title-main">{esc(b["exam_title"])}</span>'
                f'<span class="header-title-date">{esc(b["date_note"])}</span>'
                '</div>'
                f'<div class="header-info-row">{info_spans}</div>'
                f'<div class="header-name-row">{esc(b["office_field"])}</div>'
                f'<div class="header-name-row">{esc(b["name_field"])}</div>'
                '</div>'
            )
        elif t == "model_title":
            parts.append(f'<h3 class="model-title">{esc(b["text"])}</h3>')
        elif t == "section_heading":
            color = b.get("color", "#111111")
            parts.append(f'<p class="section-heading" style="color:{esc(color)};">{esc(b["text"])}</p>')
        elif t == "question":
            parts.append(f'<p class="question-line"><b>س {b["number"]}:</b> {esc(b["text"])}</p>')
        elif t == "note":
            parts.append(f'<p class="section-heading">{esc(b["text"])}</p>')
        elif t == "paragraph":
            parts.append(f'<p class="normal-line">{esc(b["text"])}</p>')
        elif t == "dotted_line":
            label = b.get("label")
            prefix_text = (b.get("prefix_text") or "").strip()
            if label and prefix_text:
                parts.append(f'<p class="dotted-line"><b>{esc(label)}</b> {esc(prefix_text)} {DOTS_LONG}</p>')
            elif label:
                parts.append(f'<p class="dotted-line"><b>{esc(label)}</b> {DOTS_LONG}</p>')
            else:
                parts.append(f'<p class="dotted-line">{DOTS_LONG}</p>')
        elif t == "meaning_grid":
            cells = b["cells"]
            rows = []
            for i in range(0, len(cells), 2):
                right = cells[i] if i < len(cells) else ""
                left = cells[i + 1] if i + 1 < len(cells) else ""
                if left:
                    rows.append(
                        '<div class="meaning-row">'
                        f'<div class="meaning-cell">{esc(right)}</div>'
                        f'<div class="meaning-cell">{esc(left)}</div>'
                        '</div>'
                    )
                else:
                    rows.append(f'<div class="meaning-row"><div class="meaning-cell-full">{esc(right)}</div></div>')
            parts.append('<div class="meaning-grid">' + "".join(rows) + '</div>')
    return "\n".join(parts)


def render_full_print_html(models, show_toolbar=True):
    """يبني مستند HTML كاملاً مستقلاً (A4، RTL، جاهز للطباعة): نموذجان في كل صفحة."""
    font_css = get_embedded_font_css()
    pages_html = []
    for page_models in group_models_for_pages(models):
        model_blocks_html = []
        for i, model in enumerate(page_models):
            if i > 0:
                model_blocks_html.append('<div class="model-divider"></div>')
            inner = render_blocks_to_html(model["blocks"])
            model_blocks_html.append(f'<div class="model-block">{inner}</div>')
        pages_html.append(f'<section class="exam-page">{"".join(model_blocks_html)}</section>')
    body = "\n".join(pages_html)

    toolbar_html = (
        '<div class="toolbar no-print"><button onclick="window.print()">'
        '🖨️ طباعة / حفظ PDF</button></div>' if show_toolbar else ""
    )

    return f"""<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>ورقة الاختبار</title>
<style>
{font_css}
:root {{
    --page-w: {PAGE_WIDTH_MM}mm;
    --page-h: {PAGE_HEIGHT_MM}mm;
    --margin: {PAGE_MARGIN_MM}mm;
}}
* {{ box-sizing: border-box; }}
html, body {{
    margin: 0; padding: 0; background: #e5e5e5;
    font-family: 'ExamArabicFont', 'Traditional Arabic', 'Amiri', Tahoma, Arial, sans-serif;
}}
.toolbar {{
    text-align: center; padding: 10px; background: #fff;
    position: sticky; top: 0; z-index: 10; border-bottom: 1px solid #ccc;
}}
.toolbar button {{
    padding: 8px 22px; font-size: 15px; cursor: pointer; border-radius: 6px;
    border: 1px solid #16306e; background: #1E3A8A; color: #fff;
}}
.pages-wrapper {{
    display: flex; flex-direction: column; align-items: center;
    gap: 10mm; padding: 10mm 0; overflow-x: auto;
    -webkit-overflow-scrolling: touch;
}}
.exam-page {{
    width: var(--page-w);
    min-height: var(--page-h);
    padding: var(--margin);
    background: #fff;
    direction: rtl;
    text-align: right;
    box-shadow: 0 0 6px rgba(0,0,0,0.25);
    flex-shrink: 0;
}}
.exam-page:not(:last-child) {{ page-break-after: always; }}
.model-block {{
    page-break-inside: avoid;
    break-inside: avoid;
}}
.model-divider {{
    border-top: 1px dashed #999;
    margin: 5px 0 6px 0;
}}
.header-box {{ border: 1.5px solid #000; border-radius: 4px; padding: 4px 8px; margin-bottom: 5px; background: #fafafa; }}
.header-title-row {{
    display: flex; flex-wrap: wrap; justify-content: space-between; align-items: baseline;
    gap: 3px 14px; font-weight: bold; font-size: 12pt; margin-bottom: 2px;
}}
.header-title-main {{ flex: 0 1 auto; }}
.header-title-date {{ flex: 0 1 auto; white-space: nowrap; }}
.header-info-row {{
    display: flex; flex-wrap: wrap; justify-content: flex-end; gap: 3px 14px;
    font-size: 9.5pt;
}}
.info-field {{ white-space: nowrap; }}
.header-name-row {{
    font-size: 9.5pt; margin-top: 2px; word-break: break-all;
}}
.model-title {{ text-align: center; border-bottom: 1px solid #ccc; padding-bottom: 2px; margin: 2px 0 5px 0; font-size: 11.5pt; }}
.section-heading {{ font-weight: bold; font-size: 11pt; margin: 5px 0 2px 0; }}
.question-line {{ font-size: 10pt; margin: 1.5px 0; line-height: 1.35; }}
.normal-line {{ font-size: 9.5pt; margin: 1.5px 0; }}
.dotted-line {{ font-size: 9pt; color: #444; margin: 1.5px 0; word-break: break-all; line-height: 1.35; }}
.meaning-grid {{ font-size: 9.5pt; margin-top: 1.5px; }}
.meaning-row {{ display: flex; justify-content: space-between; gap: 14px; margin: 3px 0; }}
.meaning-cell {{ flex: 1; word-break: break-all; }}
.meaning-cell-full {{ flex: 1; word-break: break-all; }}
@page {{ size: A4; margin: 0; }}
@media print {{
    .no-print, .toolbar {{ display: none !important; }}
    html, body {{ background: #fff; }}
    .pages-wrapper {{ padding: 0; gap: 0; overflow: visible; }}
    .exam-page {{ box-shadow: none; width: {PAGE_WIDTH_MM}mm; min-height: {PAGE_HEIGHT_MM}mm; }}
}}
@media (max-width: 480px) {{
    .exam-page {{ width: var(--page-w); }}
}}
</style>
</head>
<body>
{toolbar_html}
<div class="pages-wrapper">
{body}
</div>
</body>
</html>"""
# ==================================================================
# تقسيم النماذج إلى صفحات (نموذجان كاملان في كل صفحة A4 واحدة) -
# تستخدمه كل من صفحة الطباعة (HTML) وملف Word معاً.
# ==================================================================

def group_models_for_pages(models, per_page=MODELS_PER_PAGE):
    """يقسّم قائمة النماذج إلى مجموعات صفحات (كل مجموعة = نماذج صفحة A4 واحدة)."""
    pages = []
    for i in range(0, len(models), per_page):
        pages.append(models[i:i + per_page])
    return pages
# ==================================================================
# تصدير Word (المخرج النهائي الأساسي للتطبيق - مستند قابل للتعديل
# الكامل، عربي RTL). يقرأ نفس قائمة blocks المستخدمة في صفحة HTML —
# لا يُعاد توليد أي محتوى بشكل منفصل هنا، والنص يبقى نصاً حقيقياً
# قابلاً للتعديل (وليس صورة/لقطة شاشة).
# ==================================================================

class DocxNotAvailableError(Exception):
    pass

DOCX_FONT_NAME = "Arial"
# خط "Traditional Arabic" مخصَّص لنص الآيات القرآنية (الرسم العثماني)
# تحديداً: يعرض علامات التشكيل المركّبة والحروف الخاصة (مثل الألف
# الوصل ٱ والألف الخنجرية) بشكل أوضح من الخطوط الافتراضية مثل Arial،
# وهو مثبَّت افتراضياً مع أي إصدار من Microsoft Office يدعم العربية.
# إن لم يكن متوفراً على جهاز المستخدم، يستبدله Word تلقائياً بأقرب
# خط عربي متاح دون أي عطل.
QURAN_FONT_NAME = "Traditional Arabic"


def _docx_set_paragraph_rtl(paragraph, align_right=True):
    if align_right:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # المستند الافتراضي في python-docx يفرض تباعد أسطر 1.15x ومسافة
    # "بعد" 10pt على كل فقرة عبر docDefaults؛ نُلغي ذلك صراحةً هنا حتى
    # تنجح كل قياسات التباعد المحسوبة يدوياً (مطلوب لإتاحة نموذجين في
    # نفس صفحة A4).
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def _docx_set_run_rtl(run, size=12, bold=False, color=None, font_name=DOCX_FONT_NAME):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rpr.append(rtl)
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _docx_paragraph(doc, text, size=12, bold=False, align_center=False, color=None, space_after=6, font_name=None):
    p = doc.add_paragraph()
    _docx_set_paragraph_rtl(p, align_right=not align_center)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    _docx_set_run_rtl(run, size=size, bold=bold, color=color, font_name=font_name or DOCX_FONT_NAME)
    return p


def _zero_cell_paragraph_spacing(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def _set_table_cell_margins(table, top=20, bottom=20, left=60, right=60):
    """يقلّص هوامش خلايا الجدول الافتراضية في Word (تكون كبيرة نسبياً)
    لتوفير مساحة رأسية كافية لعرض نموذجين كاملين في صفحة A4 واحدة.
    القيم بوحدة twips (1/20 نقطة)."""
    tblPr = table._tbl.tblPr
    tblCellMar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tblCellMar.append(el)
    tblPr.append(tblCellMar)


def _hex_to_rgb_docx(hex_color, default=None):
    if not hex_color:
        return default
    try:
        h = hex_color.lstrip("#")
        return tuple(int(h[i:i + 2], 16) for i in (0, 2, 4))
    except Exception:
        return default


def _docx_set_document_page(doc):
    section = doc.sections[0]
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)
    section.page_height = Cm(PAGE_HEIGHT_MM / 10)
    section.page_width = Cm(PAGE_WIDTH_MM / 10)
    section.left_margin = Cm(PAGE_MARGIN_MM / 10)
    section.right_margin = Cm(PAGE_MARGIN_MM / 10)
    section.top_margin = Cm(PAGE_MARGIN_MM / 10)
    section.bottom_margin = Cm(PAGE_MARGIN_MM / 10)


def usable_width_for_doc(doc):
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def _set_table_fixed_layout(table):
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _set_table_rtl_visual(table):
    """
    يضبط اتجاه أعمدة الجدول بصرياً من اليمين لليسار (w:bidiVisual).
    بدون هذا الضبط الصريح، تعرض Microsoft Word الأصلية أعمدة أي جدول
    بترتيب من اليسار لليمين افتراضياً بصرف النظر عن اتجاه الفقرات
    داخلها (محاذاة/اتجاه كل فقرة على حدة لا يكفي لعكس ترتيب الأعمدة
    نفسها) - وهذا يشمل كل الجداول التي يبنيها هذا الملف: جدول
    الترويسة، جدول معاني الكلمات، وجداول أسطر الإجابة.
    """
    tblPr = table._tbl.tblPr
    bidi_visual = OxmlElement("w:bidiVisual")
    tblPr.append(bidi_visual)


def _add_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def _set_cell_border_edge(cell, edge="bottom", val="dotted", sz="8", color="595959", space="1"):
    """
    يضبط حافة واحدة محدَّدة (غالباً bottom) لخلية جدول لتُستخدم كـ"سطر
    إجابة" مستقر بصرياً في Word (بديل Word-أصلي عن تكرار نقاط "." التي
    قد تلتف بشكل غير متوقع). هذا يتجاوز حدود الجدول العامة لهذه
    الخلية تحديداً فقط، دون التأثير على باقي الجدول.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    tcBorders.append(el)


def _style_label_cell(cell, text, size=9.5, color=(90, 90, 90), bold=False):
    """يملأ خلية بنص تسمية قصير، محاذى لليمين (تُستخدم مع bidiVisual)."""
    p = cell.paragraphs[0]
    _docx_set_paragraph_rtl(p, align_right=True)
    _zero_cell_paragraph_spacing(p)
    _docx_set_run_rtl(p.add_run(text), size=size, bold=bold, color=color)


def _style_blank_line_cell(cell):
    """يجعل خلية فارغة تُستخدم كسطر إجابة عبر حد سفلي منقّط بدلاً من
    نقاط نصية متكررة."""
    p = cell.paragraphs[0]
    _docx_set_paragraph_rtl(p, align_right=False)
    _zero_cell_paragraph_spacing(p)
    p.add_run("")
    _set_cell_border_edge(cell, edge="bottom", val="dotted", sz="8", color="595959")


def _docx_render_header_box(doc, block, usable_width):
    label_width = int(usable_width * 0.22)
    line_width = usable_width - label_width

    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    # مع تفعيل bidiVisual أدناه، يصبح العمود رقم 0 هو الأقصى يميناً
    # (والعمود الأخير هو الأقصى يساراً) - لذلك: عمود 0 = التسمية
    # (تظهر يميناً كما تُقرأ العربية)، عمود 1 = سطر الإجابة الفارغ.
    table.columns[0].width = label_width
    table.columns[1].width = line_width
    for row in table.rows:
        for cell in row.cells:
            _docx_set_paragraph_rtl(cell.paragraphs[0], align_right=False)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _zero_cell_paragraph_spacing(cell.paragraphs[0])

    # الصف 1: عنوان الامتحان + التاريخ (خلية واحدة ممتدة على العمودين)
    title_cell = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    title_line = f"{block['exam_title']}        {block['date_note']}"
    title_run = title_cell.paragraphs[0].add_run(title_line)
    _docx_set_run_rtl(title_run, size=11.5, bold=True)

    # الصف 2: الفرقة | المستوى | رقم الحلقة (خلية واحدة ممتدة)
    info_cell = table.rows[1].cells[0].merge(table.rows[1].cells[1])
    info_line = "   |   ".join(block["info_fields"])
    info_run = info_cell.paragraphs[0].add_run(info_line)
    _docx_set_run_rtl(info_run, size=10, bold=False)

    # الصف 3: مكتب — تسمية + سطر إجابة (حد سفلي منقّط) بدلاً من نقاط نصية
    office_label = block["office_field"].split(".", 1)[0].rstrip()
    _docx_render_label_and_line_row(table.rows[2], office_label, usable_width)

    # الصف 4: اسم الطالب — نفس أسلوب سطر الإجابة
    name_label = block["name_field"].split(".", 1)[0].rstrip()
    _docx_render_label_and_line_row(table.rows[3], name_label, usable_width)

    _add_table_borders(table)
    _set_table_fixed_layout(table)
    _set_table_rtl_visual(table)
    _set_table_cell_margins(table, top=15, bottom=15, left=60, right=60)


def _docx_render_label_and_line_row(row, label_text, usable_width):
    """يملأ صفاً في جدول الترويسة بخليتين: تسمية (يمين، عمود 0 مع
    bidiVisual) وسطر إجابة فارغ بحد سفلي منقّط (يسار، عمود 1) بدلاً
    من الاعتماد على نقاط نصية."""
    label_cell, line_cell = row.cells[0], row.cells[1]
    _style_label_cell(label_cell, label_text, size=10, color=None)
    _style_blank_line_cell(line_cell)
    _set_cell_border_edge(line_cell, edge="bottom", val="dotted", sz="8", color="595959")


def _docx_render_meaning_grid(doc, block, usable_width):
    """
    يعرض شبكة معاني الكلمات (رقم + سطر إجابة) بنفس أسلوب الحد السفلي
    المنقّط: كل صف يحتوي عنصرين (رقم1 + سطر1 + رقم2 + سطر2) في جدول
    من 4 أعمدة، أو عنصراً واحداً ممتداً في الصف الأخير الفردي.
    """
    cells = block["cells"]
    rows_needed = (len(cells) + 1) // 2

    num_width = int(usable_width * 0.05)
    line_width = int(usable_width / 2) - num_width

    table = doc.add_table(rows=rows_needed, cols=4)
    table.autofit = False
    # ترتيب الأعمدة مع bidiVisual (0=الأقصى يميناً ... 3=الأقصى يساراً):
    # رقم١ | سطر١ | رقم٢ | سطر٢
    table.columns[0].width = num_width
    table.columns[1].width = line_width
    table.columns[2].width = num_width
    table.columns[3].width = usable_width - num_width - line_width - num_width
    _set_table_fixed_layout(table)
    _set_table_rtl_visual(table)
    _set_table_cell_margins(table, top=8, bottom=8, left=10, right=10)

    for row_idx in range(rows_needed):
        right_full = cells[row_idx * 2] if row_idx * 2 < len(cells) else ""
        left_full = cells[row_idx * 2 + 1] if row_idx * 2 + 1 < len(cells) else ""
        row = table.rows[row_idx]

        right_label = right_full.split(".", 1)[0].rstrip() if right_full else ""

        if not left_full:
            # عنصر أخير فردي: تسمية في العمود 0 وسطر إجابة ممتد عبر
            # بقية الأعمدة الثلاثة
            line_cell = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
            _style_label_cell(row.cells[0], right_label, size=10, color=None)
            _style_blank_line_cell(line_cell)
        else:
            left_label = left_full.split(".", 1)[0].rstrip()
            _style_label_cell(row.cells[0], right_label, size=10, color=None)
            _style_blank_line_cell(row.cells[1])
            _style_label_cell(row.cells[2], left_label, size=10, color=None)
            _style_blank_line_cell(row.cells[3])


def _docx_render_block(doc, block, usable_width):
    t = block["type"]

    if t == "header_box":
        _docx_render_header_box(doc, block, usable_width)

    elif t == "model_title":
        _docx_paragraph(doc, block["text"], size=11.5, bold=True, align_center=True, space_after=4)

    elif t == "section_heading":
        color = _hex_to_rgb_docx(block.get("color"))
        _docx_paragraph(doc, block["text"], size=11, bold=True, color=color, space_after=1.5)

    elif t == "question":
        _docx_paragraph(doc, f"س {block['number']}: {block['text']}", size=10.5, space_after=1.5, font_name=QURAN_FONT_NAME)

    elif t == "note":
        _docx_paragraph(doc, block["text"], size=11, bold=True, space_after=1.5)

    elif t == "paragraph":
        _docx_paragraph(doc, block["text"], size=10, space_after=1.5)

    elif t == "dotted_line":
        _docx_render_dotted_line(doc, block, usable_width)

    elif t == "meaning_grid":
        _docx_render_meaning_grid(doc, block, usable_width)


def _docx_render_dotted_line(doc, block, usable_width):
    """
    يعرض سطر إجابة (فقه/عقيدة/تفسير...) كجدول من عمودين: تسمية قصيرة
    + خلية فارغة بحد سفلي منقّط تمثّل مساحة الكتابة. هذا أكثر استقراراً
    في Word من الاعتماد على سلسلة طويلة من نقاط "." النصية (لا يلتف
    بشكل غير متوقع، ولا يختفي، ويبقى قابلاً للتعديل).
    """
    label = block.get("label") or ""
    prefix_text = (block.get("prefix_text") or "").strip()
    label_text = f"{label} {prefix_text}".strip() if prefix_text else label

    # مساحة أكبر للتسمية عند وجود نص مُدخَل من المستخدم (سؤال التفسير)
    label_ratio = 0.55 if prefix_text else 0.16
    label_width = int(usable_width * label_ratio)
    line_width = usable_width - label_width

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # عمود 0 = التسمية (تظهر يميناً مع bidiVisual)، عمود 1 = سطر
    # الإجابة الفارغ (يظهر يساراً).
    table.columns[0].width = label_width
    table.columns[1].width = line_width
    _set_table_fixed_layout(table)
    _set_table_rtl_visual(table)
    _set_table_cell_margins(table, top=5, bottom=5, left=15, right=15)

    label_cell, line_cell = table.rows[0].cells[0], table.rows[0].cells[1]
    _style_label_cell(label_cell, label_text, size=9.5, color=(90, 90, 90))
    _style_blank_line_cell(line_cell)


def _docx_render_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _docx_render_model(doc, model, usable_width):
    for block in model["blocks"]:
        _docx_render_block(doc, block, usable_width)


def build_docx_bytes(models):
    """يبني مستند Word كامل قابل للتعديل: نموذجان كاملان في كل صفحة."""
    if not DOCX_AVAILABLE:
        raise DocxNotAvailableError(
            "مكتبة python-docx غير مثبّتة. أضيفي 'python-docx' إلى requirements.txt."
        )
    doc = Document()
    _docx_set_document_page(doc)

    style = doc.styles["Normal"]
    style.font.name = DOCX_FONT_NAME
    style.font.size = Pt(10.5)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    usable_width = usable_width_for_doc(doc)
    pages = group_models_for_pages(models)
    for page_index, page_models in enumerate(pages):
        if page_index > 0:
            doc.add_page_break()
        for i, model in enumerate(page_models):
            if i > 0:
                _docx_render_divider(doc)
            _docx_render_model(doc, model, usable_width)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
# ==================================================================
# تحسينات بسيطة لواجهة Streamlit نفسها (الشريط الجانبي وأزرار
# التحميل) - لا علاقة لها بتصميم ورقة الاختبار نفسها، فتلك تُعرض
# داخل مكوّن HTML مستقل عبر render_full_print_html أعلاه.
# ==================================================================
st.markdown(
    """
    <style>
    div[data-testid="stDownloadButton"] button, div[data-testid="stButton"] button {
        width: 100%;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("صانع اختبارات القرآن الكريم")

if not DOCX_AVAILABLE:
    st.sidebar.error(
        "⚠️ مكتبة python-docx غير مثبّتة على الخادم، وبالتالي لا يمكن إنشاء "
        "ملف Word (المخرج الأساسي للتطبيق). أضيفي 'python-docx' إلى "
        "requirements.txt وأعيدي تثبيت المتطلبات."
    )

if not QURAN_DB_READY:
    st.sidebar.error(
        f"⚠️ تعذّر تحميل قاعدة بيانات القرآن الكريم الموثوقة: {QURAN_DB_ERROR} "
        "لا يمكن توليد أسئلة القرآن الكريم قبل إصلاح هذا. تأكدي من وجود ملف "
        "quran-uthmani.txt داخل مجلد 'data' بجانب app.py."
    )

st.sidebar.header("⚙️ إعدادات وتصميم الورقة")
exam_type = st.sidebar.selectbox("اختر نوع الاختبار:", ["الدوري الأول", "الدوري الثاني", "نصف العام", "آخر العام"])
selected_suwar = st.sidebar.multiselect("اختر سور الاختبار:", list(suwar_database.keys()), default=[])

# تحديث: اختيار عدد النماذج
num_models = st.sidebar.number_input("أدخل عدد النماذج المطلوبة:", min_value=1, max_value=50, value=1)

grade_option = st.sidebar.selectbox("حدد الفرقة الدراسية:", ["الأولى", "الثانية", "الثالثة", "الرابعة", "الخامسة", "السادسة", "السابعة", "الثامنة", "التاسعة", "العاشرة"])
level_option = st.sidebar.radio("حدد المستوى الحالي:", ["تمهيدي", "صغار", "متوسط", "كبار"])

st.sidebar.subheader("📋 تفعيل المحاور")
include_tarbawy = st.sidebar.checkbox("تفعيل محور (التربوي)", value=True)
include_tafsir = st.sidebar.checkbox("تفعيل محور (التفسير الإجمالي)", value=True)
custom_tafsir_verse = st.sidebar.text_input("آية/سؤال التفسير (اختياري):", value="")
include_meaning = st.sidebar.checkbox("تفعيل محور (معاني الكلمات)", value=True)

st.sidebar.subheader("📊 عدد مقاطع التسميع")
q_start_count = st.sidebar.slider("أسئلة البداية:", 0, 5, 1)
q_middle_count = st.sidebar.slider("أسئلة الوسط:", 0, 5, 1)

generate_button = st.sidebar.button("✨ توليد الورقة")

# ==================================================================
# التوليد + تخزين النتيجة في session_state حتى لا تختفي عند تنزيل
# ملف Word (استريملت يعيد تشغيل الصفحة عند أي تفاعل)
# ==================================================================

if generate_button:
    if not QURAN_DB_READY and (q_start_count > 0 or q_middle_count > 0):
        st.error(
            "⚠️ لا يمكن توليد أسئلة القرآن الكريم لأن قاعدة بيانات القرآن "
            "الموثوقة غير محمَّلة على الخادم (راجعي التحذير في الشريط الجانبي)."
        )
        st.session_state.pop("generated_models", None)
    elif not selected_suwar and (q_start_count > 0 or q_middle_count > 0):
        st.error("⚠️ يرجى اختيار سورة واحدة على الأقل من القائمة الجانبية قبل توليد الورقة.")
        st.session_state.pop("generated_models", None)
    else:
        run_seed = random.randint(1, 10_000_000)
        raw_models, general_warnings = generate_all_models(
            num_models=int(num_models),
            q_start_count=q_start_count,
            q_middle_count=q_middle_count,
            selected_suwar=selected_suwar,
            suwar_database=suwar_database,
            seed=run_seed,
        )

        built_models = []
        for i, m in enumerate(raw_models, start=1):
            built_models.append(build_exam_model_data(
                model_index=i,
                exam_type=exam_type,
                grade_option=grade_option,
                level_option=level_option,
                questions=m["questions"],
                include_tarbawy=include_tarbawy,
                include_tafsir=include_tafsir,
                custom_tafsir_verse=custom_tafsir_verse,
                include_meaning=include_meaning,
                warnings=m["warnings"],
            ))

        st.session_state["generated_models"] = built_models
        st.session_state["generated_general_warnings"] = general_warnings

# ==================================================================
# عرض النتيجة (إن وُجدت): نفس تصميم صفحة A4 يظهر في المعاينة، ويُستخدم
# بالحرف الواحد في زر الطباعة وفي تصدير Word (كلاهما يقرأ نفس بيانات
# "blocks" لكل نموذج - مصدر واحد للحقيقة)
# ==================================================================

if "generated_models" in st.session_state:
    models = st.session_state["generated_models"]
    general_warnings = st.session_state.get("generated_general_warnings", [])

    for w in general_warnings:
        st.error(w)

    per_model_warnings = [(m["model_index"], w) for m in models for w in m["warnings"]]
    if per_model_warnings:
        with st.expander("⚠️ ملاحظات حول تنوّع الأسئلة بين النماذج", expanded=False):
            for idx, w in per_model_warnings:
                st.warning(f"نموذج {idx}: {w}")

    # -- زر التحميل الأساسي: ملف Word القابل للتعديل بالكامل --------
    try:
        docx_bytes = build_docx_bytes(models)
        st.download_button(
            label="📄 تحميل ملف Word (المخرج النهائي القابل للتعديل)",
            data=docx_bytes,
            file_name="exam.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
        )
    except DocxNotAvailableError as e:
        st.error(str(e))
    except Exception as e:
        st.error(f"تعذّر إنشاء ملف Word: {e}")

    print_html_standalone = render_full_print_html(models, show_toolbar=True)
    with st.expander("🖨️ خيار إضافي: صفحة طباعة مباشرة من المتصفح (HTML)"):
        st.caption(
            "اختياري: نسخة HTML يمكن فتحها في أي متصفح (بما في ذلك الجوال) "
            "وطباعتها أو حفظها كـ PDF مباشرة من نظام التشغيل، دون الحاجة إلى Word."
        )
        st.download_button(
            label="⬇️ تحميل صفحة الطباعة (HTML)",
            data=print_html_standalone,
            file_name="exam_print.html",
            mime="text/html",
            use_container_width=True,
        )

    st.caption(
        "المعاينة أدناه هي نفس تصميم ورقة A4 التي يحتويها ملف Word تماماً "
        "(نفس بيانات النماذج والأسئلة). زر '🖨️ طباعة / حفظ PDF' داخل المعاينة "
        "يطبع صفحات الاختبار فقط، دون الشريط الجانبي."
    )

    # -- المعاينة: نفس مستند HTML المستخدم للطباعة، مضمّن مباشرة ---
    preview_height = min(900, 260 + 620 * min(len(models), 2))
    components.html(print_html_standalone, height=preview_height, scrolling=True)
